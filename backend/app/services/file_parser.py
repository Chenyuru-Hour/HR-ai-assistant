import zipfile

import pdfplumber
from docx import Document
from lxml import etree

from app.core.logging import get_logger

logger = get_logger(__name__)


def parse_pdf(file_path: str) -> dict:
    """提取PDF文件中的文字"""
    pages_text = []
    total_pages = 0

    with pdfplumber.open(file_path) as pdf:
        total_pages = len(pdf.pages)
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)

    full_text = "\n".join(pages_text)
    return {
        "total_pages": total_pages,
        "text": full_text,
    }


def parse_docx(file_path: str) -> dict:
    """提取Word文件中的文字（支持段落、表格、文本框）"""
    doc = Document(file_path)
    all_text = []

    # 1. 提取正文段落
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            all_text.append(text)

    # 2. 提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                all_text.append(" | ".join(row_text))

    # 3. 提取文本框内容（Word XML 中的 w:txbxContent）
    try:
        with zipfile.ZipFile(file_path) as z:
            xml_content = z.read("word/document.xml")
            tree = etree.fromstring(xml_content)
            nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

            # 查找所有文本框中的文本元素
            txbx_texts = tree.xpath("//w:txbxContent//w:t", namespaces=nsmap)

            # 将文本片段组合起来
            current_line = []
            for elem in txbx_texts:
                if elem.text:
                    current_line.append(elem.text)

            if current_line:
                # 合并所有文本框内容
                txbx_content = "".join(current_line)
                all_text.append(txbx_content)
    except Exception as exc:
        logger.warning("提取文本框内容失败: %s", exc)

    full_text = "\n".join(all_text)
    return {
        "total_pages": len(all_text) if all_text else 1,
        "text": full_text,
    }


def parse_file(file_path: str, file_type: str) -> dict:
    """根据文件类型自动选择解析器"""
    logger.info("开始解析文件：%s (类型： %s）", file_path, file_type)

    if file_type == "pdf":
        result = parse_pdf(file_path)
    elif file_type == "docx":
        result = parse_docx(file_path)
    else:
        raise ValueError(f"不支持的文件类型：{file_type}")

    logger.info("解析完成，提取文字长度：%d", len(result["text"]))
    return result
